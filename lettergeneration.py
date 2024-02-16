from docx import Document
import json
import datetime

def load_test_data(file_path):
    with open(file_path, 'r') as file:
        return json.load(file)
    except FileNotFoundError:
        print(f"Error: The file {file_path} was not found.")
        return []
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from {file_path}.")
        return []

def replace_placeholders(paragraph, student_data):
    if 'current_date' in paragraph.text:
        paragraph.text = paragraph.text.replace('current_date', datetime.datetime.now().strftime('%Y-%m-%d'))
    if 'STUDENT_NAME' in paragraph.text:
        paragraph.text = paragraph.text.replace('STUDENT_NAME', student_data["Student Name"])
    for i in range(1, 6):
        placeholder = f'x{i}'
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(student_data.get(placeholder, '')))

def replace_additional_paragraphs(paragraph, student_data):
    # Check for additional paragraph placeholders like 'par_2', 'par_3', etc.
    for key in student_data:
        if key.startswith('par_') and key in paragraph.text:
            paragraph.text = student_data[key]

def replace_specific_table_placeholders(table, student_data):
    # Placeholder patterns for question objectives, correct answers, and student answers
    obj_placeholders = [f'obj{i}x' for i in range(1, 33)]
    ac_placeholders = [f'ac{i}x' for i in range(1, 33)]
    ans_placeholders = [f'a{i}x' for i in range(1, 33)]

    for row in table.rows:
        for cell in row.cells:
            for placeholder in obj_placeholders + ac_placeholders + ans_placeholders:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, str(student_data.get(placeholder, '')))

def generate_test_result_letter(student_data, template_path):
    try:
        doc = Document(template_path)
    except FileNotFoundError:
        print(f"Error: The template file {template_path} was not found.")
        return

    # Replace data placeholders in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders(paragraph, student_data)
        replace_additional_paragraphs(paragraph, student_data)
    
    # Replace specific placeholders in the table
    for table in doc.tables:
        replace_specific_table_placeholders(table, student_data)

    output_filename = f'result_letter_for_{student_data.get("Student Name", "Unknown")}.docx'
    try:
        doc.save(output_filename)
    except Exception as e:
        print(f"Error saving document {output_filename}: {e}")

def main():
    test_data = load_test_data('test_data.json')
    if not test_data:
        print("No test data loaded. Exiting...")
        return

    template_path = 'student_letter_template.docx'
    for student in test_data:
        if not isinstance(student, dict):
            print("Invalid student data format. Skipping...")
            continue
        generate_test_result_letter(student, template_path)

if __name__ == "__main__":
    main()
